from docx import Document
from docx.shared import Pt
from io import BytesIO

def generate_docx(quiz_data):
    """
    Generates a Word document formatted for Microsoft Forms Quick Import.
    Includes a separate section for the answer key (Fasit).
    """
    document = Document()
    
    # Title
    document.add_heading('Quiz for MS Forms Import', 0)
    document.add_paragraph('Last opp denne filen til Microsoft Forms via "Quick Import".')
    document.add_paragraph()
    
    # --- Questions Section ---
    for i, q in enumerate(quiz_data, 1):
        # Question Text (e.g., "1. Hva er hovedstaden i Norge?")
        # MS Forms requires: Number, dot, space, question text.
        p = document.add_paragraph()
        runner = p.add_run(f"{i}. {q['question']}")
        runner.bold = True
        
        # Options (e.g., "A. Oslo")
        # MS Forms requires: Letter, dot, space, option text.
        options = q['options']
        correct_indices = q.get('correct_indices', [])
        
        for j, option in enumerate(options):
            letter = chr(65 + j) # 65 is 'A'
            p_opt = document.add_paragraph()
            runner = p_opt.add_run(f"{letter}. {option}")
            
            # Bold if correct (Visual cue for manual selection in Forms)
            if j in correct_indices:
                runner.bold = True
            
        # Add empty line for separation (Critical for MS Forms)
        document.add_paragraph()

    # --- Answer Key Section ---
    document.add_page_break()
    document.add_heading('Fasit (Kopier manuelt til MS Forms)', level=1)
    document.add_paragraph('MS Forms Quick Import støtter foreløpig ikke import av begrunnelser. Du kan kopiere dem herfra.')
    
    for i, q in enumerate(quiz_data, 1):
        document.add_heading(f"Spørsmål {i}", level=2)
        
        # Correct Answer
        # Correct Answer(s)
        correct_indices = q.get('correct_indices', [])
        
        # If for some reason it's empty, try to fallback or skip
        if not correct_indices:
            p = document.add_paragraph()
            p.add_run("Riktig svar: ").bold = True
            p.add_run("Ingen riktig svar funnet (Feil i generering)")
        else:
            correct_letters = []
            correct_texts = []
            
            for idx in correct_indices:
                if 0 <= idx < len(q['options']):
                    letter = chr(65 + idx)
                    correct_letters.append(letter)
                    correct_texts.append(f"{letter}. {q['options'][idx]}")
            
            p = document.add_paragraph()
            if len(correct_letters) > 1:
                p.add_run("Riktige svar: ").bold = True
            else:
                p.add_run("Riktig svar: ").bold = True
                
            p.add_run(", ".join(correct_texts))
        
        # Explanation
        if 'explanation' in q and q['explanation']:
            p = document.add_paragraph()
            p.add_run("Begrunnelse: ").bold = True
            p.add_run(q['explanation'])
            
        document.add_paragraph() # Spacing

    # Save to BytesIO
    f = BytesIO()
    document.save(f)
    f.seek(0)
    return f
